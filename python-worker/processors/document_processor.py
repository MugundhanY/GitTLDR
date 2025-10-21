"""
Document Processor for GitTLDR Python Worker.
Uses LangChain document loaders to handle various file formats including PDF, DOCX, XLSX, CSV, images, audio, etc.
"""
import os
import tempfile
import base64
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import mimetypes
from io import BytesIO

from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Handles document processing for various file formats using LangChain loaders."""

    def __init__(self):
        self.supported_formats = {
            # Text formats
            '.txt': 'text',
            '.md': 'text',
            '.markdown': 'text',
            '.json': 'json',
            '.xml': 'xml',
            '.html': 'html',
            '.htm': 'html',
            '.css': 'text',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.py': 'python',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.cs': 'csharp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala',
            '.sh': 'shell',
            '.bash': 'shell',
            '.sql': 'sql',
            '.r': 'r',
            '.yaml': 'yaml',
            '.yml': 'yaml',

            # Document formats
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.xlsx': 'xlsx',
            '.xls': 'xls',
            '.pptx': 'pptx',
            '.ppt': 'ppt',

            # Data formats
            '.csv': 'csv',

            # Image formats
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.gif': 'image',
            '.bmp': 'image',
            '.tiff': 'image',
            '.tif': 'image',
            '.webp': 'image',

            # Audio formats
            '.mp3': 'audio',
            '.wav': 'audio',
            '.flac': 'audio',
            '.m4a': 'audio',
            '.ogg': 'audio',
        }

        # MIME type mappings for better detection
        self.mime_mappings = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'application/vnd.ms-excel': 'xls',
            'text/csv': 'csv',
            'text/plain': 'text',
            'text/markdown': 'text',
            'application/json': 'json',
            'text/html': 'html',
            'text/xml': 'xml',
            'image/jpeg': 'image',
            'image/png': 'image',
            'image/gif': 'image',
            'image/bmp': 'image',
            'image/tiff': 'image',
            'audio/mpeg': 'audio',
            'audio/wav': 'audio',
            'audio/flac': 'audio',
            'audio/mp4': 'audio',
            'audio/ogg': 'audio',
        }

    def detect_format(self, filename: str, mime_type: str = None) -> str:
        """Detect the file format from filename and/or MIME type."""
        # First try MIME type
        if mime_type and mime_type in self.mime_mappings:
            return self.mime_mappings[mime_type]

        # Then try file extension
        if filename:
            _, ext = os.path.splitext(filename.lower())
            if ext in self.supported_formats:
                return self.supported_formats[ext]

        # Try to guess MIME type from filename
        if filename:
            guessed_mime, _ = mimetypes.guess_type(filename)
            if guessed_mime and guessed_mime in self.mime_mappings:
                return self.mime_mappings[guessed_mime]

        return 'unknown'

    async def process_document(self, content: Any, filename: str, mime_type: str = None) -> Dict[str, Any]:
        """
        Process a document using appropriate LangChain loader based on format.

        Args:
            content: Document content (bytes, string, or base64)
            filename: Original filename
            mime_type: MIME type if known

        Returns:
            Dict with processed content and metadata
        """
        try:
            # Detect format
            format_type = self.detect_format(filename, mime_type)
            logger.info(f"Processing document: {filename} (format: {format_type}, mime: {mime_type})")

            # Convert content to bytes if needed
            content_bytes = self._normalize_content(content)

            if not content_bytes:
                return {
                    'success': False,
                    'error': 'No content provided',
                    'content': '',
                    'metadata': {'format': format_type, 'filename': filename}
                }

            # Process based on format
            if format_type == 'pdf':
                result = await self._process_pdf(content_bytes, filename)
            elif format_type == 'docx':
                result = await self._process_docx(content_bytes, filename)
            elif format_type in ['xlsx', 'xls']:
                result = await self._process_excel(content_bytes, filename)
            elif format_type == 'csv':
                result = await self._process_csv(content_bytes, filename)
            elif format_type == 'image':
                result = await self._process_image(content_bytes, filename)
            elif format_type == 'audio':
                result = await self._process_audio(content_bytes, filename)
            elif format_type == 'json':
                result = await self._process_json(content_bytes, filename)
            else:
                # Default text processing
                result = await self._process_text(content_bytes, filename)

            result['metadata'] = result.get('metadata', {})
            result['metadata'].update({
                'format': format_type,
                'filename': filename,
                'mime_type': mime_type,
                'original_size': len(content_bytes)
            })

            return result

        except Exception as e:
            logger.error(f"Failed to process document {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': '',
                'metadata': {'format': format_type, 'filename': filename}
            }

    def _normalize_content(self, content: Any) -> Optional[bytes]:
        """Normalize content to bytes."""
        if content is None:
            return None

        if isinstance(content, bytes):
            return content
        elif isinstance(content, str):
            # Try base64 decode first
            try:
                return base64.b64decode(content)
            except Exception:
                # If not base64, encode as UTF-8
                return content.encode('utf-8')
        else:
            # Try to convert to string then bytes
            try:
                return str(content).encode('utf-8')
            except Exception:
                return None

    async def _process_pdf(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF document."""
        try:
            # Try multiple PDF libraries in order of preference
            text_content = None
            metadata = {}

            # Try PyMuPDF (fitz) first - best for text extraction
            try:
                import fitz
                doc = fitz.open(stream=content_bytes, filetype="pdf")
                text_content = ""
                page_count = len(doc)  # Get page count before processing
                for page in doc:
                    text_content += page.get_text() + "\n\n"
                metadata = {
                    'pages': page_count,
                    'library': 'PyMuPDF'
                }
                doc.close()
                logger.info(f"Extracted text from PDF using PyMuPDF: {len(text_content)} chars, {page_count} pages")
            except ImportError:
                logger.warning("PyMuPDF not available, trying pdfplumber")

            # Try pdfplumber as fallback
            if not text_content or len(text_content.strip()) < 100:
                try:
                    import pdfplumber
                    with pdfplumber.open(BytesIO(content_bytes)) as pdf:
                        text_content = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_content += page_text + "\n\n"
                        metadata = {
                            'pages': len(pdf.pages),
                            'library': 'pdfplumber'
                        }
                    logger.info(f"Extracted text from PDF using pdfplumber: {len(text_content)} chars, {len(pdf.pages)} pages")
                except ImportError:
                    logger.warning("pdfplumber not available, trying PyPDF2")

            # Try PyPDF2 as last resort
            if not text_content or len(text_content.strip()) < 100:
                try:
                    from PyPDF2 import PdfReader
                    pdf_file = BytesIO(content_bytes)
                    pdf_reader = PdfReader(pdf_file)
                    text_content = ""
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n\n"
                    metadata = {
                        'pages': len(pdf_reader.pages),
                        'library': 'PyPDF2'
                    }
                    logger.info(f"Extracted text from PDF using PyPDF2: {len(text_content)} chars, {len(pdf_reader.pages)} pages")
                except ImportError:
                    logger.warning("PyPDF2 not available")

            if text_content and len(text_content.strip()) > 0:
                return {
                    'success': True,
                    'content': text_content.strip(),
                    'metadata': metadata
                }
            else:
                return {
                    'success': False,
                    'error': 'No text could be extracted from PDF (may contain only images)',
                    'content': f'[PDF Document - {filename}]\nNo text content could be extracted. This PDF may contain only images or scanned pages.',
                    'metadata': {'pages': 'unknown', 'library': 'none'}
                }

        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[PDF Document - {filename}]\nError processing PDF: {str(e)}'
            }

    async def _process_docx(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX document."""
        try:
            from docx import Document
            doc = Document(BytesIO(content_bytes))

            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                text_content += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
                text_content += "[END TABLE]\n\n"

            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }

            return {
                'success': True,
                'content': text_content.strip(),
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[DOCX Document - {filename}]\nError processing DOCX: {str(e)}'
            }

    async def _process_excel(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process Excel document."""
        try:
            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(BytesIO(content_bytes))
            all_content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Convert to markdown table
                table_content = f"\n## Sheet: {sheet_name}\n\n"
                table_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"

                # Add column info
                columns_info = f"Columns: {', '.join(df.columns.tolist())}\n\n"
                table_content += columns_info

                # Add data preview (first 50 rows)
                if len(df) > 0:
                    preview_df = df.head(50)
                    table_content += preview_df.to_markdown(index=False)
                    if len(df) > 50:
                        table_content += f"\n\n... ({len(df) - 50} more rows)"
                else:
                    table_content += "(Empty sheet)"

                all_content.append(table_content)

            metadata = {
                'sheets': len(excel_file.sheet_names),
                'sheet_names': excel_file.sheet_names
            }

            return {
                'success': True,
                'content': "\n\n".join(all_content),
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"Excel processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[Excel Document - {filename}]\nError processing Excel file: {str(e)}'
            }

    async def _process_csv(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV document."""
        try:
            import pandas as pd

            # Try different encodings and separators
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            separators = [',', ';', '\t', '|']
            df = None
            used_encoding = None
            used_separator = None

            # First, try to auto-detect with multiple encoding/separator combinations
            for encoding in encodings:
                for sep in separators:
                    try:
                        df = pd.read_csv(
                            BytesIO(content_bytes), 
                            encoding=encoding,
                            sep=sep,
                            on_bad_lines='skip',  # Skip problematic lines
                            engine='python'  # More flexible parser
                        )
                        
                        # CRITICAL: Validate that parsing was actually successful
                        # Reject results with:
                        # 1. Too many columns (>1000 suggests wrong delimiter)
                        # 2. Zero rows with many columns (header parsed as data)
                        # 3. Garbage column names (non-printable characters)
                        is_valid = True
                        
                        # Check 1: Reasonable column count (most CSVs have < 100 columns)
                        if df.shape[1] > 1000:
                            logger.debug(f"Rejecting {encoding}/{sep}: too many columns ({df.shape[1]})")
                            is_valid = False
                        
                        # Check 2: Must have either multiple columns OR some rows
                        elif not (df.shape[1] > 1 or len(df) > 0):
                            logger.debug(f"Rejecting {encoding}/{sep}: no data")
                            is_valid = False
                        
                        # Check 3: If 0 rows but many columns, likely wrong delimiter
                        elif len(df) == 0 and df.shape[1] > 50:
                            logger.debug(f"Rejecting {encoding}/{sep}: 0 rows with {df.shape[1]} columns (wrong delimiter)")
                            is_valid = False
                        
                        # Check 4: Column names should be mostly readable text
                        elif df.shape[1] > 0:
                            first_col_name = str(df.columns[0])
                            # Count non-printable/non-ASCII characters
                            non_printable = sum(1 for c in first_col_name if ord(c) > 127 or not c.isprintable())
                            if len(first_col_name) > 0 and non_printable / len(first_col_name) > 0.5:
                                logger.debug(f"Rejecting {encoding}/{sep}: garbage column names")
                                is_valid = False
                        
                        if is_valid:
                            used_encoding = encoding
                            used_separator = sep
                            logger.info(f"Successfully parsed CSV with encoding={encoding}, separator='{sep}', shape={df.shape}")
                            break
                        else:
                            df = None  # Reset for next attempt
                            
                    except Exception as e:
                        continue
                
                if df is not None and used_separator is not None:
                    break

            # If still failed, try with error handling
            if df is None or (df.shape[1] == 1 and len(df) == 0):
                logger.warning(f"Standard CSV parsing failed, trying with more permissive settings")
                try:
                    # Try with python engine and more permissive settings
                    auto_df = pd.read_csv(
                        BytesIO(content_bytes),
                        encoding='utf-8',
                        sep=None,  # Auto-detect separator
                        engine='python',
                        on_bad_lines='skip',
                        encoding_errors='ignore'  # Ignore encoding errors
                    )
                    
                    # CRITICAL: Validate auto-detected result with same criteria
                    is_valid_auto = True
                    
                    # Check 1: Reasonable column count
                    if auto_df.shape[1] > 1000:
                        logger.warning(f"Auto-detect rejected: too many columns ({auto_df.shape[1]})")
                        is_valid_auto = False
                    
                    # Check 2: Must have data
                    elif not (auto_df.shape[1] > 1 or len(auto_df) > 0):
                        logger.warning(f"Auto-detect rejected: no data")
                        is_valid_auto = False
                    
                    # Check 3: 0 rows with many columns = wrong delimiter
                    elif len(auto_df) == 0 and auto_df.shape[1] > 50:
                        logger.warning(f"Auto-detect rejected: 0 rows with {auto_df.shape[1]} columns (wrong delimiter)")
                        is_valid_auto = False
                    
                    # Check 4: Column names should be readable
                    elif auto_df.shape[1] > 0:
                        first_col_name = str(auto_df.columns[0])
                        non_printable = sum(1 for c in first_col_name if ord(c) > 127 or not c.isprintable())
                        if len(first_col_name) > 0 and non_printable / len(first_col_name) > 0.5:
                            logger.warning(f"Auto-detect rejected: garbage column names ({first_col_name[:50]}...)")
                            is_valid_auto = False
                    
                    if is_valid_auto:
                        df = auto_df
                        used_encoding = 'utf-8 (auto)'
                        used_separator = 'auto-detected'
                        logger.info(f"Auto-detect successful: shape={df.shape}")
                    else:
                        # Auto-detect also failed validation
                        logger.error(f"Auto-detect produced invalid result, falling back to raw text")
                        df = None
                        
                except Exception as e:
                    logger.warning(f"Auto-detect parsing failed: {e}")
                    df = None
                
                # Last resort: try to read as plain text
                if df is None:
                    try:
                        text_content = content_bytes.decode('utf-8', errors='ignore')
                        return {
                            'success': False,
                            'error': 'Could not parse as CSV - showing raw content',
                            'content': f'[CSV Document - {filename}]\n\nCould not parse as structured CSV. Raw content preview:\n\n{text_content[:5000]}\n\n... (truncated if longer)',
                            'metadata': {'raw_size': len(content_bytes)}
                        }
                    except Exception:
                        raise ValueError("Could not decode CSV with any supported method")

            if df is None or len(df.columns) == 0:
                raise ValueError("CSV parsing resulted in empty dataframe")

            # Format as markdown
            content = f"## CSV Data: {filename}\n\n"
            content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
            if used_encoding and used_separator:
                content += f"Encoding: {used_encoding}, Separator: '{used_separator}'\n\n"

            # Column information
            content += "### Columns:\n"
            for i, col in enumerate(df.columns):
                dtype = str(df[col].dtype)
                # Get sample values
                non_null_values = df[col].dropna()
                sample = non_null_values.head(3).tolist() if len(non_null_values) > 0 else []
                sample_str = f" (examples: {sample})" if sample else ""
                content += f"{i+1}. {col} ({dtype}){sample_str}\n"
            content += "\n"

            # Data preview
            content += "### Data Preview:\n\n"
            preview_rows = min(50, len(df))
            preview_df = df.head(preview_rows)
            
            # Use to_string for better formatting if to_markdown fails
            try:
                content += preview_df.to_markdown(index=False)
            except Exception:
                content += preview_df.to_string(index=False)

            if len(df) > preview_rows:
                content += f"\n\n... ({len(df) - preview_rows} more rows)"

            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'dtypes': {col: str(df[col].dtype) for col in df.columns},
                'encoding': used_encoding,
                'separator': used_separator
            }

            return {
                'success': True,
                'content': content,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"CSV processing failed: {str(e)}")
            # Try to at least show raw content
            try:
                text_content = content_bytes.decode('utf-8', errors='ignore')
                return {
                    'success': False,
                    'error': str(e),
                    'content': f'[CSV Document - {filename}]\nError processing CSV: {str(e)}\n\nRaw content preview:\n{text_content[:2000]}'
                }
            except Exception:
                return {
                    'success': False,
                    'error': str(e),
                    'content': f'[CSV Document - {filename}]\nError processing CSV: {str(e)}'
                }

    async def _process_image(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process image document."""
        try:
            from PIL import Image

            # Open image
            image = Image.open(BytesIO(content_bytes))

            # Get basic image info
            width, height = image.size
            format_name = image.format or "Unknown"

            # Build base content
            content = f"[Image: {filename}]\n"
            content += f"Format: {format_name}\n"
            content += f"Dimensions: {width} × {height} pixels\n"
            content += f"File size: {len(content_bytes)} bytes\n\n"

            # Try OCR with tesseract if available
            ocr_text = ""
            ocr_success = False
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(image)
                ocr_success = len(ocr_text.strip()) > 0
                if ocr_success:
                    content += f"OCR Text Extracted:\n{ocr_text.strip()}\n\n"
                    logger.info(f"Successfully extracted OCR text from {filename}: {len(ocr_text)} chars")
            except ImportError:
                logger.info(f"pytesseract not available, trying Gemini Vision API for {filename}")
            except Exception as ocr_error:
                logger.warning(f"tesseract OCR failed: {str(ocr_error)}, trying Gemini Vision API")

            # If OCR failed or not available, try Gemini Vision API
            if not ocr_success:
                try:
                    # Use Gemini's vision capabilities to analyze the image
                    from services.gemini_client import gemini_client
                    import base64
                    
                    # Convert image bytes to base64
                    image_base64 = base64.b64encode(content_bytes).decode('utf-8')
                    
                    # Create vision prompt
                    vision_prompt = """Analyze this image and provide:
1. A detailed description of what you see
2. Any text visible in the image (OCR)
3. Key visual elements, colors, and composition
4. Purpose or context of the image if discernible

Be thorough but concise."""
                    
                    # Call Gemini with vision
                    vision_result = await gemini_client.analyze_image_with_vision(
                        image_base64=image_base64,
                        prompt=vision_prompt,
                        mime_type=f"image/{format_name.lower()}" if format_name != "Unknown" else "image/png"
                    )
                    
                    if vision_result and len(vision_result.strip()) > 0:
                        content += f"AI Vision Analysis:\n{vision_result.strip()}\n\n"
                        logger.info(f"Successfully analyzed image with Gemini Vision: {len(vision_result)} chars")
                        ocr_success = True
                    else:
                        content += "No text or meaningful content could be extracted from this image.\n"
                        
                except Exception as vision_error:
                    logger.warning(f"Gemini Vision API failed: {str(vision_error)}")
                    content += f"Could not extract text or analyze image content.\n"
                    content += f"(OCR not available, Vision API error: {str(vision_error)})\n"

            # If still no text extracted
            if not ocr_success:
                content += "No text could be extracted from this image using OCR or AI vision."

            metadata = {
                'width': width,
                'height': height,
                'format': format_name,
                'has_ocr': ocr_success,
                'ocr_text_length': len(ocr_text) if ocr_text else 0,
                'extraction_method': 'tesseract' if ocr_text else 'gemini_vision' if ocr_success else 'none'
            }

            return {
                'success': True,
                'content': content,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"Image processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[Image: {filename}]\nError processing image: {str(e)}'
            }

    async def _process_audio(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process audio document."""
        try:
            # Try faster-whisper first (better performance)
            transcript = None
            metadata = {}

            try:
                import faster_whisper
                model = faster_whisper.WhisperModel("base")  # Small model for speed

                # Save to temporary file for processing
                with tempfile.NamedTemporaryFile(suffix=os.path.splitext(filename)[1], delete=False) as temp_file:
                    temp_file.write(content_bytes)
                    temp_path = temp_file.name

                try:
                    segments, info = model.transcribe(temp_path, language="en")
                    transcript = " ".join([segment.text for segment in segments])

                    metadata = {
                        'duration': info.duration,
                        'language': info.language,
                        'library': 'faster-whisper'
                    }

                    logger.info(f"Transcribed audio using faster-whisper: {len(transcript)} chars, duration: {info.duration}s")

                finally:
                    os.unlink(temp_path)

            except ImportError:
                logger.warning("faster-whisper not available, trying openai-whisper")

            # Fallback to openai-whisper
            if not transcript:
                try:
                    import whisper

                    # Save to temporary file
                    with tempfile.NamedTemporaryFile(suffix=os.path.splitext(filename)[1], delete=False) as temp_file:
                        temp_file.write(content_bytes)
                        temp_path = temp_file.name

                    try:
                        model = whisper.load_model("base")
                        result = model.transcribe(temp_path)

                        transcript = result["text"]
                        metadata = {
                            'duration': result.get('duration', 0),
                            'language': result.get('language', 'unknown'),
                            'library': 'openai-whisper'
                        }

                        logger.info(f"Transcribed audio using openai-whisper: {len(transcript)} chars")

                    finally:
                        os.unlink(temp_path)

                except ImportError:
                    logger.warning("openai-whisper not available")

            if transcript and len(transcript.strip()) > 0:
                content = f"[Audio Transcript: {filename}]\n\n{transcript.strip()}"
                return {
                    'success': True,
                    'content': content,
                    'metadata': metadata
                }
            else:
                return {
                    'success': False,
                    'error': 'Could not transcribe audio',
                    'content': f'[Audio File: {filename}]\nCould not transcribe audio content.',
                    'metadata': {'library': 'none'}
                }

        except Exception as e:
            logger.error(f"Audio processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[Audio File: {filename}]\nError processing audio: {str(e)}'
            }

    async def _process_json(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process JSON document."""
        try:
            import json

            # Try to parse JSON
            json_str = content_bytes.decode('utf-8')
            json_data = json.loads(json_str)

            # Format for display
            content = f"[JSON Document: {filename}]\n\n"
            content += json.dumps(json_data, indent=2)

            metadata = {
                'is_array': isinstance(json_data, list),
                'keys_count': len(json_data) if isinstance(json_data, dict) else len(json_data) if isinstance(json_data, list) else 0
            }

            return {
                'success': True,
                'content': content,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"JSON processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[JSON Document: {filename}]\nError processing JSON: {str(e)}'
            }

    async def _process_text(self, content_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = None

            for encoding in encodings:
                try:
                    text_content = content_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                raise ValueError("Could not decode text with any supported encoding")

            metadata = {
                'encoding': encoding,
                'lines': len(text_content.split('\n')),
                'characters': len(text_content)
            }

            return {
                'success': True,
                'content': text_content,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"Text processing failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f'[Text Document: {filename}]\nError processing text: {str(e)}'
            }


# Global instance
document_processor = DocumentProcessor()